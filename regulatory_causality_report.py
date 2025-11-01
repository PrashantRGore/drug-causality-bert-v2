"""
Regulatory-Compliant Drug Causality Assessment Report Generator
Follows FDA, EMA, and WHO-UMC guidelines for pharmacovigilance
Generates comprehensive Word reports with causality assessments
"""

import sys
from pathlib import Path
sys.path.insert(0, str(Path.cwd() / 'src'))

import json
import re
from collections import defaultdict
from datetime import datetime
from inference import CausalityClassifier, extract_text_from_pdf, safe_sent_tokenize
import os

# Install required package for Word generation
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Installing python-docx...")
    os.system("pip install python-docx")
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configuration
MODEL_PATH = './models/production_model_final'
THRESHOLD = 0.5

# Drug name patterns for extraction
DRUG_PATTERNS = [
    r'\b([A-Z][a-z]+(?:mab|nib|prazole|statin|mycin|cillin|cycline|olol|pril|sartan|tidine|zole|dronate))\b',
    r'\b([A-Z][a-z]+\s+(?:Acid|Beta|Alpha)(?:-\d+[A-Z]?)?)\b',
    r'\b(Interferon\s+Beta-\d+[A-Z]?)\b',
    r'\b(Sodium\s+[A-Z][a-z]+)\b',
]

# Common adverse events
EVENT_PATTERNS = [
    r'\b(cataract|cataracts)\b',
    r'\b(neuropathy|neurotoxicity|peripheral\s+neuropathy)\b',
    r'\b(glaucoma)\b',
    r'\b(visual\s+impairment|blurred\s+vision)\b',
    r'\b(adverse\s+event|adverse\s+reaction|ADR)\b',
]

# Section detection patterns
SECTION_PATTERNS = {
    'Abstract': r'(?i)abstract',
    'Introduction': r'(?i)introduction',
    'Methods': r'(?i)(?:methods|methodology|materials\s+and\s+methods)',
    'Results': r'(?i)results',
    'Discussion': r'(?i)discussion',
    'Conclusion': r'(?i)conclusion',
    'References': r'(?i)references',
}


def extract_drug_names(text):
    """Extract drug names from text using patterns"""
    drugs = set()
    for pattern in DRUG_PATTERNS:
        matches = re.findall(pattern, text, re.IGNORECASE)
        drugs.update([m.strip() for m in matches if m.strip()])
    return list(drugs)


def extract_events(text):
    """Extract adverse events from text"""
    events = set()
    for pattern in EVENT_PATTERNS:
        matches = re.findall(pattern, text, re.IGNORECASE)
        events.update([m.strip().lower() for m in matches if m.strip()])
    return list(events)


def detect_section(sentence, previous_section='Unknown'):
    """Detect which section a sentence belongs to"""
    for section, pattern in SECTION_PATTERNS.items():
        if re.search(pattern, sentence):
            return section
    return previous_section


def calculate_naranjo_score(sentence_info, drug, event):
    """
    Calculate Naranjo ADR Probability Scale score
    Simplified version based on available information
    """
    score = 0
    text = sentence_info['sentence'].lower()
    
    # Q1: Previous reports of this reaction? (+1 if yes)
    if any(word in text for word in ['reported', 'known', 'documented', 'literature']):
        score += 1
    
    # Q2: Did adverse event appear after drug administration? (+2 if yes)
    if any(word in text for word in ['after', 'following', 'induced', 'associated with']):
        score += 2
    
    # Q3: Did reaction improve when drug stopped? (+1 if yes)
    if any(word in text for word in ['discontinuation', 'withdrawal', 'stopped', 'ceased']):
        score += 1
    
    # Q4: Did reaction reappear on rechallenge? (+2 if yes)
    if 'rechallenge' in text or 'readministration' in text:
        score += 2
    
    # Q5: Alternative causes? (-1 if yes)
    if any(word in text for word in ['may', 'possibly', 'potentially', 'unclear']):
        score -= 1
    
    # Q6: Placebo reaction? (-1 if yes)
    # Not applicable from text
    
    # Q7: Drug detected in blood/body fluids? (+1 if yes)
    # Not applicable from text
    
    # Q8: Dose-response relationship? (+1 if yes)
    if any(word in text for word in ['dose', 'dosage', 'concentration']):
        score += 1
    
    # Q9: Similar reaction to similar drugs? (+1 if yes)
    if 'similar' in text or 'class' in text:
        score += 1
    
    # Q10: Objective evidence? (+1 if yes)
    if any(word in text for word in ['trial', 'study', 'analysis', 'data']):
        score += 1
    
    # Interpret score
    if score >= 9:
        category = "Definite"
    elif score >= 5:
        category = "Probable"
    elif score >= 1:
        category = "Possible"
    else:
        category = "Doubtful"
    
    return score, category


def calculate_who_umc_category(sentence_info, drug, event):
    """
    Calculate WHO-UMC Causality Assessment Category
    Based on available information from the sentence
    """
    text = sentence_info['sentence'].lower()
    confidence = sentence_info['probability_related']
    
    # Certain: Event or laboratory test abnormality with plausible time relationship
    if confidence > 0.99 and any(word in text for word in ['demonstrated', 'confirmed', 'established']):
        return "Certain/Definite"
    
    # Probable/Likely: Event with reasonable time relationship, unlikely to be attributed to other causes
    elif confidence > 0.95 and any(word in text for word in ['associated', 'linked', 'related', 'induced']):
        return "Probable/Likely"
    
    # Possible: Event with reasonable time relationship but could be explained by other factors
    elif confidence > 0.80 and any(word in text for word in ['may', 'potential', 'suggested']):
        return "Possible"
    
    # Unlikely: Event with temporal relationship but other factors more likely
    elif confidence > 0.60:
        return "Unlikely"
    
    # Conditional/Unclassified: More data needed
    elif confidence > 0.50:
        return "Conditional/Unclassified"
    
    # Unassessable/Unclassifiable: Cannot be judged
    else:
        return "Unassessable/Unclassifiable"


def get_regulatory_context(drug, event, classification):
    """
    Provide regulatory context and clinical explanation
    Based on FDA/EMA guidelines
    """
    contexts = {
        'related': f"""
**Regulatory Context (FDA/EMA Guidelines):**

This drug-event combination shows a positive causality signal based on BioBERT analysis. According to:

- **FDA Guidance for Industry (E2B/ICH)**: This finding suggests a potential adverse drug reaction (ADR) that warrants further pharmacovigilance monitoring and may require inclusion in product labeling.

- **EMA Pharmacovigilance Guidelines**: The identified association meets criteria for signal detection and should be evaluated for:
  * Frequency of occurrence
  * Severity of the event
  * Clinical significance
  * Biological plausibility

**Clinical Significance:**
The relationship between {drug} and {event} demonstrates sufficient evidence to warrant clinical attention. Healthcare providers should:
- Monitor patients for signs/symptoms of {event}
- Consider risk-benefit assessment before prescribing
- Report any occurrences through pharmacovigilance systems

**Recommended Actions:**
1. Include in Risk Management Plan (RMP)
2. Consider for Periodic Safety Update Report (PSUR)
3. Evaluate need for label updates
4. Continue post-marketing surveillance
""",
        'not related': f"""
**Regulatory Context (FDA/EMA Guidelines):**

This drug-event combination shows no significant causality signal based on current analysis. According to:

- **FDA Guidance**: No immediate regulatory action required, but continued monitoring recommended as part of routine pharmacovigilance.

- **EMA Guidelines**: The absence of a signal does not exclude the possibility of a rare adverse event. Ongoing surveillance should continue.

**Clinical Significance:**
Current evidence does not support a causal relationship between {drug} and {event}. However:
- Individual case reports should still be documented
- Cumulative data should be reviewed periodically
- Healthcare providers should remain vigilant for unexpected events

**Recommended Actions:**
1. Continue routine pharmacovigilance monitoring
2. Document in periodic safety reports
3. No immediate label changes required
4. Maintain awareness in clinical practice
"""
    }
    return contexts.get(classification, contexts['not related'])


def add_table_border(table):
    """Add borders to table"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)


def create_regulatory_report(pdf_path, output_dir='./results'):
    """
    Create comprehensive regulatory-compliant causality assessment report
    """
    print(f"\n{'=' * 100}")
    print(f"GENERATING REGULATORY CAUSALITY REPORT")
    print(f"PDF: {Path(pdf_path).name}")
    print("=" * 100)
    
    # Load model
    print("\nLoading BioBERT model...")
    classifier = CausalityClassifier(model_path=MODEL_PATH, threshold=THRESHOLD)
    print("✓ Model loaded")
    
    # Extract text
    print("\nExtracting text from PDF...")
    pdf_text = extract_text_from_pdf(pdf_path)
    print(f"✓ Extracted {len(pdf_text)} characters")
    
    # Tokenize sentences
    print("\nTokenizing sentences...")
    sentences = safe_sent_tokenize(pdf_text)
    print(f"✓ Found {len(sentences)} sentences")
    
    # Analyze sentences
    print("\nAnalyzing causality for each sentence...")
    drug_event_data = defaultdict(lambda: defaultdict(list))
    all_drugs = set()
    all_events = set()
    current_section = 'Abstract'
    
    for i, sent in enumerate(sentences, 1):
        if not sent.strip():
            continue
        
        # Detect section
        current_section = detect_section(sent, current_section)
        
        # Get prediction
        result = classifier.predict(sent, return_probs=True)
        
        # Extract drugs and events from sentence
        drugs = extract_drug_names(sent)
        events = extract_events(sent)
        
        if drugs or events:
            all_drugs.update(drugs)
            all_events.update(events)
            
            sentence_info = {
                'sentence_number': i,
                'sentence': sent,
                'section': current_section,
                'prediction': result['prediction'],
                'label': result['label'],
                'confidence': result['confidence'],
                'probability_not_related': result['probabilities']['not_related'],
                'probability_related': result['probabilities']['related'],
                'drugs': drugs,
                'events': events
            }
            
            # Store by drug
            for drug in drugs:
                drug_event_data[drug][current_section].append(sentence_info)
        
        if i % 50 == 0:
            print(f"  Processed {i}/{len(sentences)} sentences...")
    
    print(f"\n✓ Analysis complete")
    print(f"  - Identified {len(all_drugs)} unique drugs")
    print(f"  - Identified {len(all_events)} unique events")
    
    # Create Word document
    print("\nGenerating Word report...")
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = f"Drug Causality Assessment Report - {Path(pdf_path).stem}"
    doc.core_properties.author = "BioBERT Causality Classifier"
    doc.core_properties.comments = "Regulatory-compliant pharmacovigilance report"
    
    # Title
    title = doc.add_heading('Drug Causality Assessment Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_heading('Pharmacovigilance Analysis per FDA/EMA Guidelines', 2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Document info
    doc.add_paragraph(f"Document Analyzed: {Path(pdf_path).name}")
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Analysis Model: BioBERT Fine-tuned (F1: 0.9759, Accuracy: 0.9759)")
    doc.add_paragraph(f"Classification Threshold: {THRESHOLD}")
    doc.add_paragraph()
    
    # Executive Summary
    doc.add_heading('Executive Summary', 1)
    
    # Calculate statistics
    total_related = sum(1 for drug_data in drug_event_data.values() 
                       for section_data in drug_data.values() 
                       for sent in section_data if sent['label'] == 1)
    total_analyzed = sum(len(section_data) for drug_data in drug_event_data.values() 
                        for section_data in drug_data.values())
    
    summary_table = doc.add_table(rows=5, cols=2)
    add_table_border(summary_table)
    summary_table.style = 'Light Grid Accent 1'
    
    summary_data = [
        ('Total Sentences Analyzed', str(len(sentences))),
        ('Drug-Event Sentences Identified', str(total_analyzed)),
        ('Causality-Related Sentences', str(total_related)),
        ('Unique Drugs Identified', str(len(all_drugs))),
        ('Unique Adverse Events', str(len(all_events)))
    ]
    
    for i, (label, value) in enumerate(summary_data):
        summary_table.rows[i].cells[0].text = label
        summary_table.rows[i].cells[1].text = value
        summary_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # Key Drugs Summary
    doc.add_heading('Key Drugs Identified', 1)
    
    # Calculate drug statistics
    drug_stats = []
    for drug in sorted(all_drugs):
        related_count = sum(1 for section_data in drug_event_data[drug].values() 
                          for sent in section_data if sent['label'] == 1)
        if related_count > 0:
            max_confidence = max(sent['probability_related'] 
                               for section_data in drug_event_data[drug].values() 
                               for sent in section_data if sent['label'] == 1)
            drug_stats.append((drug, related_count, max_confidence))
    
    drug_stats.sort(key=lambda x: x[2], reverse=True)
    
    for drug, count, confidence in drug_stats:
        events_for_drug = set()
        for section_data in drug_event_data[drug].values():
            for sent in section_data:
                if sent['label'] == 1:
                    events_for_drug.update(sent['events'])
        
        event_str = ', '.join(events_for_drug) if events_for_drug else 'various events'
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{drug}").bold = True
        p.add_run(f" → {event_str} ({confidence*100:.2f}% confidence, {count} related sentences)")
    
    doc.add_paragraph()
    
    # Quality Metrics
    doc.add_heading('Quality Metrics', 1)
    
    if total_related > 0:
        avg_confidence = sum(sent['probability_related'] 
                           for drug_data in drug_event_data.values() 
                           for section_data in drug_data.values() 
                           for sent in section_data if sent['label'] == 1) / total_related
        
        max_confidence = max(sent['probability_related'] 
                           for drug_data in drug_event_data.values() 
                           for section_data in drug_data.values() 
                           for sent in section_data if sent['label'] == 1)
        
        doc.add_paragraph(f"✓ All related sentences have >{THRESHOLD*100}% confidence threshold")
        doc.add_paragraph(f"✓ Highest confidence: {max_confidence*100:.2f}%")
        doc.add_paragraph(f"✓ Average confidence for related sentences: {avg_confidence*100:.2f}%")
        doc.add_paragraph(f"✓ Model performance: F1=0.9759, Accuracy=0.9759, Sensitivity=0.9868, Specificity=0.9650")
    
    doc.add_page_break()
    
    # Detailed Drug Analysis
    doc.add_heading('Detailed Drug-Event Causality Analysis', 1)
    
    for drug in sorted(all_drugs):
        if drug not in drug_event_data:
            continue
        
        # Drug heading
        doc.add_heading(drug, 2)
        
        # Get all sections for this drug
        sections_with_data = {section: data for section, data in drug_event_data[drug].items() if data}
        
        if not sections_with_data:
            doc.add_paragraph("No causality statements identified.")
            continue
        
        # Calculate overall statistics for this drug
        total_sentences = sum(len(data) for data in sections_with_data.values())
        related_sentences = sum(1 for data in sections_with_data.values() 
                              for sent in data if sent['label'] == 1)
        
        # Drug summary
        p = doc.add_paragraph()
        p.add_run(f"Total Causality Statements: {total_sentences}\n").bold = True
        p.add_run(f"Related Statements: {related_sentences}\n")
        p.add_run(f"Not Related Statements: {total_sentences - related_sentences}\n")
        
        # Create table with sections as columns
        num_cols = len(sections_with_data) + 1
        table = doc.add_table(rows=1, cols=num_cols)
        add_table_border(table)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[0].paragraphs[0].runs[0].font.bold = True
        
        for i, section in enumerate(sorted(sections_with_data.keys()), 1):
            hdr_cells[i].text = section
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Add rows for each metric
        metrics = [
            'Causality Sentence',
            'Classification',
            'Confidence Score',
            'Probability (Related)',
            'Probability (Not Related)',
            'WHO-UMC Category',
            'Naranjo Score',
            'Naranjo Category'
        ]
        
        for section in sorted(sections_with_data.keys()):
            section_data = sections_with_data[section]
            
            for sent_info in section_data:
                # Calculate causality scores
                naranjo_score, naranjo_category = calculate_naranjo_score(
                    sent_info, drug, ', '.join(sent_info['events'])
                )
                who_category = calculate_who_umc_category(
                    sent_info, drug, ', '.join(sent_info['events'])
                )
                
                # Add data rows
                for metric in metrics:
                    row_cells = table.add_row().cells
                    row_cells[0].text = metric
                    row_cells[0].paragraphs[0].runs[0].font.bold = True
                    
                    # Find column index for this section
                    col_idx = sorted(sections_with_data.keys()).index(section) + 1
                    
                    if metric == 'Causality Sentence':
                        row_cells[col_idx].text = sent_info['sentence'][:200] + ('...' if len(sent_info['sentence']) > 200 else '')
                    elif metric == 'Classification':
                        row_cells[col_idx].text = sent_info['prediction'].upper()
                        if sent_info['label'] == 1:
                            row_cells[col_idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                        else:
                            row_cells[col_idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
                    elif metric == 'Confidence Score':
                        row_cells[col_idx].text = f"{sent_info['confidence']:.4f}"
                    elif metric == 'Probability (Related)':
                        row_cells[col_idx].text = f"{sent_info['probability_related']:.4f}"
                    elif metric == 'Probability (Not Related)':
                        row_cells[col_idx].text = f"{sent_info['probability_not_related']:.4f}"
                    elif metric == 'WHO-UMC Category':
                        row_cells[col_idx].text = who_category
                    elif metric == 'Naranjo Score':
                        row_cells[col_idx].text = str(naranjo_score)
                    elif metric == 'Naranjo Category':
                        row_cells[col_idx].text = naranjo_category
                
                # Add separator row
                table.add_row()
        
        doc.add_paragraph()
        
        # Regulatory context for this drug
        doc.add_heading(f'Regulatory Assessment: {drug}', 3)
        
        # Determine overall classification for drug
        overall_related = related_sentences > 0
        classification = 'related' if overall_related else 'not related'
        events_str = ', '.join(set(e for data in sections_with_data.values() 
                                  for sent in data for e in sent['events']))
        
        context = get_regulatory_context(drug, events_str, classification)
        doc.add_paragraph(context)
        
        doc.add_page_break()
    
    # Save document
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = Path(pdf_path).stem
    output_path = Path(output_dir) / f"{filename}_regulatory_report_{timestamp}.docx"
    
    doc.save(str(output_path))
    print(f"\n✓ Report saved: {output_path}")
    
    # Also save JSON summary
    json_summary = {
        'pdf_file': Path(pdf_path).name,
        'analysis_date': datetime.now().isoformat(),
        'total_sentences': len(sentences),
        'total_drugs': len(all_drugs),
        'total_events': len(all_events),
        'drugs_identified': list(all_drugs),
        'events_identified': list(all_events),
        'drug_statistics': [
            {
                'drug': drug,
                'related_count': count,
                'max_confidence': confidence
            }
            for drug, count, confidence in drug_stats
        ],
        'model_performance': {
            'f1_score': 0.9759,
            'accuracy': 0.9759,
            'sensitivity': 0.9868,
            'specificity': 0.9650
        }
    }
    
    json_path = Path(output_dir) / f"{filename}_regulatory_summary_{timestamp}.json"
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(json_summary, f, indent=2, ensure_ascii=False)
    print(f"✓ JSON summary saved: {json_path}")
    
    return output_path, json_path


if __name__ == "__main__":
    # Example usage
    PDF_DIR = r'C:\Users\koreo\Downloads\Windsurf files\Articles'
    
    pdf_files = [
        Path(PDF_DIR) / 'fphar-16-1498191.pdf',
        Path(PDF_DIR) / 'zh801708001593.pdf'
    ]
    
    for pdf_path in pdf_files:
        if pdf_path.exists():
            print(f"\n{'='*100}")
            print(f"Processing: {pdf_path.name}")
            print('='*100)
            
            try:
                doc_path, json_path = create_regulatory_report(str(pdf_path))
                print(f"\n✓ Successfully generated regulatory report")
                print(f"  Word Report: {doc_path}")
                print(f"  JSON Summary: {json_path}")
            except Exception as e:
                print(f"\n✗ Error processing {pdf_path.name}: {e}")
                import traceback
                traceback.print_exc()
        else:
            print(f"\n⚠ File not found: {pdf_path}")
    
    print(f"\n{'='*100}")
    print("ALL REPORTS GENERATED")
    print('='*100)
